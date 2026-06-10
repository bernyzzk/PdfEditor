from pathlib import Path

import fitz
from PySide6.QtWidgets import QApplication
from PySide6.QtGui import QImage, QColor

from pdf_editor.model import PdfDocument


def make_pdf(path: Path) -> None:
    doc = fitz.open()
    for label in ("Première page", "Deuxième page", "Troisième page"):
        page = doc.new_page()
        page.insert_text((72, 72), label)
    doc.save(path)
    doc.close()


def test_editing_workflow(tmp_path: Path) -> None:
    QApplication.instance() or QApplication([])
    source = tmp_path / "source.pdf"
    output = tmp_path / "output.pdf"
    make_pdf(source)

    model = PdfDocument()
    assert model.open(str(source))
    assert model.page_count == 3
    assert not model.render_page(0).isNull()

    model.add_text(0, fitz.Rect(72, 100, 300, 130), "Texte ajouté", 12, (0, 0, 0))
    objects = model.page_objects(0)
    selected = model.object_at(0, fitz.Point(80, 110))
    assert any(obj.kind == "text" for obj in objects)
    assert selected and selected.kind == "text"
    model.replace_text(0, selected, "Texte modifié", selected.size, selected.color, selected.font)
    model.add_ink(0, [fitz.Point(72, 130), fitz.Point(150, 150)], (1, 0, 0), 2)
    model.add_highlight(0, fitz.Rect(65, 55, 180, 80))
    image_path = tmp_path / "image.png"
    image = QImage(40, 30, QImage.Format.Format_RGB32)
    image.fill(QColor("#246bfd"))
    assert image.save(str(image_path))
    model.add_image(0, fitz.Rect(320, 80, 400, 140), str(image_path))
    image_obj = model.object_at(0, fitz.Point(350, 100))
    assert image_obj and image_obj.kind == "image"
    model.replace_image(0, image_obj, str(image_path))
    image_payload = model.copy_object_payload(0, image_obj)
    assert model.paste_object(0, image_payload, 120)
    moved_image_rect = fitz.Rect(350, 160, 450, 240)
    image_obj = model.object_at(0, fitz.Point(350, 100))
    model.transform_object(0, image_obj, moved_image_rect)
    moved_images = [obj for obj in model.page_objects(0) if obj.kind == "image"]
    assert any(obj.rect.intersects(moved_image_rect) for obj in moved_images)
    model.add_link(0, fitz.Rect(50, 180, 180, 205), "https://example.com")
    assert model.doc.load_page(0).get_links()
    assert model.redact_text("Première") == 1
    extra = tmp_path / "extra.pdf"
    make_pdf(extra)
    model.merge_files([str(extra)])
    assert model.page_count == 6
    extracted = tmp_path / "extracted.pdf"
    model.extract_page(0, str(extracted))
    assert fitz.open(extracted).page_count == 1
    model.rotate_page(1)
    model.move_page(2, 0)
    model.delete_page(2)
    assert model.page_count == 5

    model.undo()
    assert model.page_count == 6
    model.redo()
    assert model.page_count == 5
    model.save(str(output))

    exported = fitz.open(output)
    assert exported.page_count == 5
    assert "Texte modifié" in "".join(page.get_text() for page in exported).replace("\xa0", " ")
    assert any(page.annots() for page in exported)
    exported.close()


def test_embedded_font_edit_and_move(tmp_path: Path) -> None:
    QApplication.instance() or QApplication([])
    source = tmp_path / "font-source.pdf"
    font_path = Path(r"C:\Windows\Fonts\calibri.ttf")
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 90), "Texte Calibri original", fontsize=15, fontname="Calibri", fontfile=str(font_path))
    doc.save(source)
    doc.close()

    model = PdfDocument()
    assert model.open(str(source))
    obj = model.object_at(0, fitz.Point(100, 85))
    assert obj and "Calibri" in obj.font and obj.font_xref
    model.replace_text(0, obj, "Texte Calibri modifié", obj.size, obj.color, obj.font)
    edited = next(item for item in model.page_objects(0) if "modifié" in item.text)
    assert "Calibri" in edited.font
    model.transform_object(0, edited, fitz.Rect(120, 150, 340, 180))
    moved = next(item for item in model.page_objects(0) if "modifié" in item.text.replace("\xa0", " "))
    assert moved.rect.intersects(fitz.Rect(120, 150, 340, 180))
    payload = model.copy_object_payload(0, moved)
    duplicate = model.paste_object(0, payload, 45)
    assert duplicate and duplicate.kind == "text"


def test_subset_font_edit_supports_french_and_euro(tmp_path: Path) -> None:
    QApplication.instance() or QApplication([])
    source = tmp_path / "subset-arial.pdf"
    output = tmp_path / "subset-arial-edited.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 90),
        "Total HT",
        fontsize=11,
        fontname="ArialMT",
        fontfile=r"C:\Windows\Fonts\arial.ttf",
    )
    doc.save(source)
    doc.close()

    model = PdfDocument()
    assert model.open(str(source))
    obj = model.object_at(0, fitz.Point(90, 85))
    assert obj and obj.kind == "text"
    replacement = "Total HT modifié : 35,96 €"
    model.replace_text(0, obj, replacement, obj.size, obj.color, obj.font)
    model.save(str(output))

    edited = fitz.open(output)
    assert replacement in edited[0].get_text().replace("\xa0", " ")
    assert any("Arial" in font[3] for font in edited[0].get_fonts(full=True))
    edited.close()


def test_ocr_and_office_exports(tmp_path: Path) -> None:
    QApplication.instance() or QApplication([])
    from PIL import Image, ImageDraw, ImageFont
    from docx import Document
    from openpyxl import load_workbook
    from pptx import Presentation

    image_path = tmp_path / "scan.png"
    image = Image.new("RGB", (1200, 500), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 64)
    draw.text((60, 160), "FACTURE TEST 2026", font=font, fill="black")
    image.save(image_path)

    source = tmp_path / "scan.pdf"
    doc = fitz.open()
    page = doc.new_page(width=600, height=250)
    page.insert_image(page.rect, filename=str(image_path))
    doc.save(source)
    doc.close()

    model = PdfDocument()
    assert model.open(str(source))
    assert not model.doc[0].get_text().strip()
    assert model.run_ocr(language="eng") == 1
    assert "FACTURE" in model.doc[0].get_text().upper()

    word = tmp_path / "result.docx"
    excel = tmp_path / "result.xlsx"
    powerpoint = tmp_path / "result.pptx"
    model.export_word(str(word))
    model.export_excel(str(excel))
    model.export_powerpoint(str(powerpoint))
    assert "FACTURE" in "\n".join(p.text for p in Document(word).paragraphs).upper()
    assert "FACTURE" in str(load_workbook(excel).active["A1"].value).upper()
    assert len(Presentation(powerpoint).slides) == 1


def test_rotated_clockify_style_pdf_is_editable(tmp_path: Path) -> None:
    QApplication.instance() or QApplication([])
    source = tmp_path / "rotated-report.pdf"
    image_path = tmp_path / "logo.png"
    image = QImage(120, 50, QImage.Format.Format_RGB32)
    image.fill(QColor("#246bfd"))
    assert image.save(str(image_path))

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_text((60, 80), "Rapport hebdomadaire", fontsize=18)
    page.insert_text((60, 130), "Total: 26:54:25", fontsize=14)
    page.insert_text((60, 180), "06:42:36", fontsize=10)
    page.insert_image(fitz.Rect(60, 220, 180, 270), filename=str(image_path))
    page.set_rotation(90)
    doc.save(source)
    doc.close()

    model = PdfDocument()
    assert model.open(str(source))
    assert model.doc[0].rotation == 0
    title = next(obj for obj in model.page_objects(0) if "Rapport" in obj.text)
    number = next(obj for obj in model.page_objects(0) if "26:54:25" in obj.text)
    logo = next(obj for obj in model.page_objects(0) if obj.kind == "image")
    center = lambda rect: fitz.Point((rect.x0 + rect.x1) / 2, (rect.y0 + rect.y1) / 2)
    assert model.object_at(0, center(title.rect)).kind == "text"
    assert model.object_at(0, center(number.rect)).kind == "text"
    assert model.object_at(0, center(logo.rect)).kind == "image"
    model.replace_text(0, number, "Total: 30:00:00", number.size, number.color, number.font)
    assert "Total: 30:00:00" in model.doc[0].get_text().replace("\xa0", " ")
